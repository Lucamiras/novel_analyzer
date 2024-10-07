from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import SyllableTokenizer, word_tokenize, sent_tokenize


class DocumentLoader:
    def __init__(self, path:str):
        self.path = path
        self.doc = Document(self.path)
        self.text = self._get_full_text()
        self.chapter_names = self._get_chapter_names()
        self.chapters = self._get_chapters()
        self.author = self._get_author()
        self.title = self._get_title()

    def _get_full_text(self):
        return [paragraph.text for paragraph in self.doc.paragraphs if paragraph.text != '']

    def _get_chapters(self):
        indices = [i for i, paragraph in enumerate(self.doc.paragraphs) if paragraph.style.name.startswith('Heading 1')]
        doc_dict = {}
        l, r = 0, 1
        while r < len(indices):
            doc_dict[self.doc.paragraphs[indices[l]].text] = [self.doc.paragraphs[i].text for i in range(indices[l]+1,indices[r]) if self.doc.paragraphs[i].text != '']
            l += 1
            r += 1
        return doc_dict

    def _get_chapter_names(self):
        return [paragraph.text for paragraph in self.doc.paragraphs if paragraph.style.name.startswith('Heading 1')]    
    
    def _get_author(self):
        return self.doc.core_properties.author
    
    def _get_title(self):
        title = self.doc.core_properties.title
        return title if title else 'No title found'
    

class TextParser:
    stopwords = {
        'english': stopwords.words('english'),
        'german': stopwords.words('german')
        }
    punctuation = ['.',',',';',':','!','?','(',')','[',']','{','}','"','\'', '»', '«', '…', '-', '_']

    def __init__(self, text:list, exclusion_list:list):
        self.text = text
        self.exclusion_list = exclusion_list
        self.full_text = self._join_text()
        self.words = self._get_words()
        self.syllables = self._get_syllables()
        self.sentences = self._get_sentences()
        self.word_dict = self._generate_word_dict()
        self.word_dict_without_stopwords = self._generate_word_dict_without_stopwords()
        self.num_chars, self.chars_no_spaces = self._get_num_characters()
        self.num_words = len(self.words)
        self.num_syllables = len(self.syllables)
        self.num_sentences = len(self.sentences)
    
    def _join_text(self):
        return ' '.join(self.text)

    def _get_num_characters(self):
        return len(self.full_text), len(self.full_text.replace(' ',''))

    def _get_sentences(self):
        return sent_tokenize(self.full_text)
    
    def _get_words(self):
        return [word for word in word_tokenize(self.full_text) 
        if (word not in self.punctuation)
        and word.lower() not in self.exclusion_list]

    def _get_syllables(self):
        syl = SyllableTokenizer() 
        syllabilized_words = [syl.tokenize(word) for word in self.words]
        syllables = [syllable for word in syllabilized_words for syllable in word]
        return syllables

    def _generate_word_dict(self):
        words_lower = [word.lower() for word in self.words]
        word_dict = {}
        for word in words_lower:
            word_dict[word] = word_dict.get(word, 0) + 1
        word_dict_sorted = {
            key: value for key, value in sorted(word_dict.items(), key=lambda x: x[1], reverse=True) 
        }
        return word_dict_sorted
    
    def _generate_word_dict_without_stopwords(self):
        return {
            key: value for key, value in self.word_dict.items() if key not in self.stopwords['german']
        }
    