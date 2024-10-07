import streamlit as st
from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import SyllableTokenizer, word_tokenize, sent_tokenize
import matplotlib.pyplot as plt
import numpy as np


class DocumentAnalyzer:
    def __init__(self, path:str="", top_n:int=10):
        self.path = path
        self.top_n = top_n
        self.stopwords = {
            'english': stopwords.words('english'),
            'german': stopwords.words('german')
        }
        self.doc = Document(self.path)
        self.chapters_dict = self._get_chapters()
        self.chapters_len = {key: len(word_tokenize(' '.join(value))) for key, value in self.chapters_dict.items()}
        self.author = self.doc.core_properties.author
        self.data = [paragraph.text for paragraph in self.doc.paragraphs if paragraph.text != '']
        self.text = ' '.join(self.data)
        self.sent = sent_tokenize(self.text)
        self.word = [word for word in word_tokenize(self.text) if len(word) > 1]
        self.word_dict = self._generate_word_dict()
        self.len_text = len(self.text)
        self.len_text_without_spaces = len(self.text.replace(' ',''))
        self.num_sentences = len(self.sent)
        self.len_sentences = [len(sent) for sent in self.sent]
        self.num_words = len(self.word)
        self.len_words = [len(word) for word in self.word]
        self.num_syllables = self._get_num_syllables()
        self.lexical_richness = self._get_lexical_richness()
        self.flesch_kincaid = self._get_flesch_kincaid_score()
    
    def _generate_word_dict(self):
        words_lower = [word.lower() for word in self.word]
        word_dict = {}
        for word in words_lower:
            word_dict[word] = word_dict.get(word, 0) + 1
        word_dict_sorted = {
            key: value for key, value in sorted(word_dict.items(), key=lambda x: x[1], reverse=True) 
        }
        return word_dict_sorted
    
    def _generate_word_dict_without_stopwords(self):
        return {
            key: value for key, value in self.word_dict.items() if key not in self.stopwords['german']
        }
    
    def get_basic_stats(self):
        average_word_length = round(float(np.mean(self.len_words)),1)
        average_sent_length = round(float(np.mean(self.len_sentences)),1)
        average_para_length = round(float(np.mean([len(word_tokenize(para)) for para in self.data])))
        median_word_length = round(float(np.median(self.len_words)),1)
        median_sent_length = round(float(np.median(self.len_sentences)),1)
        median_para_length = round(float(np.median([len(word_tokenize(para)) for para in self.data])))
        return (
            average_word_length, 
            average_sent_length, 
            average_para_length, 
            median_word_length, 
            median_sent_length, 
            median_para_length
            )

    def get_top_n_most_and_least_used(self):
        word_dict_without_stopwords = self._generate_word_dict_without_stopwords()
        top_most_used = [
            key for key in word_dict_without_stopwords.items()
        ][:self.top_n]
        top_least_used = [
            key for key in word_dict_without_stopwords.items()
        ][-self.top_n:]
        return top_most_used, top_least_used
    
    def get_outliers(self):
        sentences_ordered_by_length = sorted(self.sent, key=len, reverse=True)
        words_ordered_by_length = sorted(self.word, key=len, reverse=True)
        short_sent = sentences_ordered_by_length[-1]
        long_sent = sentences_ordered_by_length[0]
        short_word = words_ordered_by_length[-1]
        long_word = words_ordered_by_length[0]
        return (
            short_sent,
            long_sent,
            short_word,
            long_word
        )

    def _get_num_syllables(self):
        syl = SyllableTokenizer() 
        syllabilized_words = [syl.tokenize(word) for word in self.word]
        syllables = [syllable for word in syllabilized_words for syllable in word]
        return len(syllables)

    def _get_lexical_richness(self):
        num_unique_words = len([value for value in self.word_dict.values() if value == 1])
        num_total_words = len(self.word_dict)
        return np.round(num_unique_words/num_total_words,2)
    
    def _get_flesch_kincaid_score(self):
        return float(
            np.round(
                206.835 - (1.015 * (self.num_words / self.num_sentences)) - (84.6 * (self.num_syllables / self.num_words)),
                1)
            )
    
    def _get_flesch_kincaid_eval(self):
        scale = [
            (range(90, 101), "5th grade", "Very easy to read. Easily understood by an average 11-year-old student."),
            (range(80, 90), "6th grade", "Easy to read. Conversational English for consumers."),
            (range(70, 80), "7th grade", "Fairly easy to read."),
            (range(60, 70), "8th & 9th grade", "Plain English. Easily understood by 13- to 15-year-old students."),
            (range(50, 60), "10th to 12th", "Fairly difficult to read."),
            (range(30, 50), "College", "Difficult to read."),
            (range(10, 30), "College graduate", "Very difficult to read. Best understood by university graduates."),
            (range(0, 10), "Professional", "Extremely difficult to read. Best understood by university graduates.")
        ]
        score = int(np.round(self.flesch_kincaid, 0))
        for sub_range, grade, description in scale:
            if score in sub_range:
                return grade, description
    
    def _get_chapters(self):
        indices = [i for i, paragraph in enumerate(self.doc.paragraphs) if paragraph.style.name.startswith('Heading 1')]
        doc_dict = {}
        l, r = 0, 1
        while r < len(indices):
            doc_dict[self.doc.paragraphs[indices[l]].text] = [self.doc.paragraphs[i].text for i in range(indices[l]+1,indices[r]) if self.doc.paragraphs[i].text != '']
            l += 1
            r += 1
        return doc_dict
    
    def plot_word_frequency(self, container, title, stopwords_removed, reverse=False):
        if stopwords_removed:
            word_dict = self._generate_word_dict_without_stopwords()
        if not stopwords_removed:
            word_dict = self.word_dict
        if reverse:
            top_words = list(word_dict.items())[-self.top_n:]    
        if not reverse:
            top_words = list(word_dict.items())[:self.top_n]
        words, frequencies = zip(*top_words)
        plt.figure(figsize=(5, 5))
        plt.bar(words, frequencies)
        plt.xlabel('Words')
        plt.ylabel('Frequencies')
        plt.title(title)
        plt.xticks(rotation=90)
        plt.ylim(0, max(list(word_dict.values()))*1.01)
        container.pyplot(plt)

    def plot_sentence_length(self, X, container, title, xlabel):
        plt.figure(figsize=(10, 2))
        plt.boxplot(X, vert=False)
        plt.xlabel(xlabel)
        plt.title(title)
        container.pyplot(plt)
        